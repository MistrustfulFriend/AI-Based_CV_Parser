from flask import Flask, request, jsonify, send_file, send_from_directory, Response
from flask_cors import CORS
from openai import OpenAI
import json
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import os
import time

app = Flask(__name__, static_folder='.')
CORS(app)

EXTRACTION_PROMPT = """Extract ALL information from this resume and return it as JSON. Keep the experience descriptions EXACTLY as written in the resume - do not modify, summarize, or rephrase anything.

Rules:
- Never change, correct, summarize, or paraphrase any text from the resume.
- Every company name, job title, task description, responsibility, project name, or educational detail must be kept exactly as it appears.
- Return the output only in the requested JSON format.
- Output valid JSON only, using double quotes for all keys and string values.
- Do not add trailing commas.
- Escape all quotes, newlines, and special characters inside strings so the JSON is parseable.
- Do not add, remove, or modify any information.
- Do not include explanations, comments, or text outside the JSON brackets.
- Copy all text exactly; do not change any wording, grammar, spelling, or capitalization.
- Do not include any bullet characters (•, -, *, etc.) from the CV or template in the extracted text.
- Do not include empty strings or empty items in arrays.

Required JSON structure:
{
  "expert": {
    "last_name": "",
    "first_name": "",
    "city_names": "",
    "nationality_en": "",
    "year_of_birth": "",
    "about": "",
    "educations": [
      {"years": "", "degree": "", "field_of_study": "", "university_name": ""}
    ],
    "courses": [
      {"year": "", "name": "", "organization": ""}
    ],
    "languages_list_en": "",
    "ai_assessment": "",
    "professional_experiences": [
      {
        "from": "",
        "to": "",
        "company": "",
        "role": "",
        "tasks": [],
        "technologies": [],
        "projects": [
          {"name": "", "role": "", "description": "", "responsibilities": ""}
        ]
      }
    ]
  }
}

CRITICAL: For professional_experiences, copy the exact text from the resume. Do not change wording, grammar, or style. Remove unnecessary line breaks, tabs, and extra spaces between words.  Correct spacing inside words that are accidentally split by spaces, so "R evising" becomes "Revising" and "Fur ther" becomes "Further".

Resume text:
"""

@app.route('/')
def home():
    return send_from_directory('.', 'index.html')

@app.route('/<path:path>')
def serve_static(path):
    return send_from_directory('.', path)

@app.route('/api/parse', methods=['POST'])
def parse_resume():
    try:
        data = request.json
        pdf_text = data.get('pdf_text')
        api_key = data.get('api_key')
        
        if not pdf_text or not api_key:
            return jsonify({'error': 'Missing required data'}), 400
        
        # Initialize OpenAI client with API key
        client = OpenAI(api_key=api_key)

        # -------------------------
        # Step 1: Extraction
        # -------------------------
        print("Step 1: Starting resume parsing...")
        
        response = client.chat.completions.create(
            model="o3-mini-2025-01-31",
            messages=[
                {"role": "system", "content": "You are a highly accurate resume parser. Your task is to extract information from resumes exactly as written. You must **never** change, correct, summarize, or paraphrase any text from the resume. Every company name, job title, task description, responsibility, project name, or educational detail must be kept **exactly as it appears**. Return the output only in the requested JSON format. Do not add, remove, or modify any information. Do not include explanations, comments, or text outside the JSON brackets. CRITICAL: Copy all text exactly, do not change any wording, grammar, spelling, or capitalization. Tasks, responsibilities, and company details must be exactly as written. Remove unnecessary line breaks, tabs, and extra spaces between words.  Correct spacing inside words that are accidentally split by spaces, so 'R evising' becomes 'Revising' and 'Fur ther' becomes 'Further'."},
                {"role": "user", "content": EXTRACTION_PROMPT + pdf_text}
            ],

            max_completion_tokens=50000
        )
        
        response_text = response.choices[0].message.content.strip()

        # Try to extract valid JSON even if the model added extra text
        try:
            parsed_data = json.loads(response_text)
        except json.JSONDecodeError as e:
            print("JSON parsing failed, trying to clean response:", e)
            start = response_text.find('{')
            end = response_text.rfind('}') + 1
            if start != -1 and end != -1:
                try:
                    parsed_data = json.loads(response_text[start:end])
                except Exception as e2:
                    print("Failed again:", e2)
                    return jsonify({'error': f'Invalid JSON from model: {response_text}'}), 500
            else:
                return jsonify({'error': f'Invalid JSON from model: {response_text}'}), 500

        print("Step 1: Resume parsing done, JSON extracted.")

        # -------------------------
        # Step 2: Verification
        # -------------------------
        print("Step 2: Starting verification...")

        verification_prompt = f"""
        Check the following resume data for consistency, missing fields, or formatting issues. Remove unnecessary line breaks, tabs, and extra spaces between words. Correct spacing inside words that are accidentally split by spaces, so "R evising" becomes "Revising" and "Fur ther" becomes "Further".
        Return a JSON with "valid": true/false and "issues": ["list of issues"].

        Resume data:
        {json.dumps(parsed_data, indent=2)}
        """

        verification_response = client.chat.completions.create(
            model="o3-mini-2025-01-31",
            messages=[
                {"role": "system", "content": "You are a data validator."},
                {"role": "user", "content": verification_prompt}
            ],
 
             max_completion_tokens=50000

        )

        validation_text = verification_response.choices[0].message.content.strip()
        try:
            validation_result = json.loads(validation_text)
        except json.JSONDecodeError as e:
            print("Verification JSON parsing failed:", e)
            # fallback: assume valid false with raw text issue
            validation_result = {"valid": False, "issues": [validation_text]}

        print("Step 2: Verification done.")

        return jsonify({
            'data': parsed_data,
            'validation': validation_result
        })
        
    except Exception as e:
        print(f"Error in parse_resume: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/download', methods=['POST'])
def generate_word():
    try:
        data = request.json
        expert = data.get('expert', {})

        print("Generating Word document...")

        doc = Document()
        
        # Title - EXPERT PROFILE (centered, bold)
        title = doc.add_heading('EXPERT PROFILE', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Create main info table (6 rows x 3 columns)
        table1 = doc.add_table(rows=6, cols=3)
        table1.style = 'Table Grid'
        
        # Set column widths
        table1.columns[0].width = Inches(1.5)
        table1.columns[1].width = Inches(3.5)
        table1.columns[2].width = Inches(1.5)
        
        # Row 0: SURNAME
        add_table_row_bold(table1.rows[0].cells[0], 'SURNAME:', expert.get('last_name', ''), table1.rows[0].cells[1])
        table1.rows[0].cells[2].text = '[Photo placeholder]'
        
        # Row 1: FIRSTNAME
        add_table_row_bold(table1.rows[1].cells[0], 'FIRSTNAME:', expert.get('first_name', ''), table1.rows[1].cells[1])
        
        # Row 2: CITY
        add_table_row_bold(table1.rows[2].cells[0], 'CITY:', expert.get('city_names', ''), table1.rows[2].cells[1])
        
        # Row 3: NATIONALITY
        add_table_row_bold(table1.rows[3].cells[0], 'NATIONALITY:', expert.get('nationality_en', ''), table1.rows[3].cells[1])
        
        # Row 4: YEAR OF BIRTH
        add_table_row_bold(table1.rows[4].cells[0], 'YEAR OF BIRTH:', expert.get('year_of_birth', ''), table1.rows[4].cells[1])
        
        # Row 5: AVAILABILITY
        add_table_row_bold(table1.rows[5].cells[0], 'AVAILABILITY:', '', table1.rows[5].cells[1])
        
        doc.add_paragraph()
        
        # Create details table
        table2 = doc.add_table(rows=10, cols=2)
        table2.style = 'Table Grid'
        table2.columns[0].width = Inches(1.8)
        table2.columns[1].width = Inches(4.7)
        
        row_idx = 0
        
        # FOCUS AREAS
        add_table_section_bold(table2.rows[row_idx].cells[0], 'FOCUS AREAS')
        table2.rows[row_idx].cells[1].text = expert.get('about', '')
        row_idx += 1
        
        # QUALIFICATION
        add_table_section_bold(table2.rows[row_idx].cells[0], 'QUALIFICATION')
        qual_cell = table2.rows[row_idx].cells[1]
        qual_para = qual_cell.paragraphs[0]
        for edu in expert.get('educations', []):
            if qual_para.text:
                qual_para = qual_cell.add_paragraph()
            qual_para.add_run(f"{edu.get('years', '')} ").bold = False
            qual_para.add_run(f"{edu.get('degree', '')}").bold = True
            qual_para.add_run(f", {edu.get('field_of_study', '')}, {edu.get('university_name', '')}")
        row_idx += 1
        
        # CERTIFICATIONS
        add_table_section_bold(table2.rows[row_idx].cells[0], 'CERTIFICATIONS')
        cert_cell = table2.rows[row_idx].cells[1]
        cert_para = cert_cell.paragraphs[0]
        for course in expert.get('courses', []):
            if cert_para.text:
                cert_para = cert_cell.add_paragraph()
            year = course.get('year', '')
            if year:
                cert_para.add_run(f"{year}: ")
            cert_para.add_run(course.get('name', '')).bold = True
            org = course.get('organization', '')
            if org:
                cert_para.add_run(f", {org}")
        row_idx += 1
        
        # LANGUAGES
        add_table_section_bold(table2.rows[row_idx].cells[0], 'LANGUAGES')
        table2.rows[row_idx].cells[1].text = expert.get('languages_list_en', '')
        row_idx += 1
        
        # HOBBIES / PRIVATE
        add_table_section_bold(table2.rows[row_idx].cells[0], 'HOBBIES / PRIVATE')
        table2.rows[row_idx].cells[1].text = ''
        row_idx += 1
        
        # SALARY / OTE EXPECTATION
        add_table_section_bold(table2.rows[row_idx].cells[0], 'SALARY / OTE EXPECTATION')
        table2.rows[row_idx].cells[1].text = ''
        row_idx += 1
        
        # SELF ASSESSMENT
        add_table_section_bold(table2.rows[row_idx].cells[0], 'SELF ASSESSMENT')
        table2.rows[row_idx].cells[1].text = ''
        row_idx += 1
        
        # OUR ASSESSMENT
        add_table_section_bold(table2.rows[row_idx].cells[0], 'OUR ASSESSMENT')
        table2.rows[row_idx].cells[1].text = expert.get('ai_assessment', '')
        row_idx += 1
        
        # YOUR CONTACT at e-aces.com
        contact_cell = table2.rows[row_idx].cells[0]
        contact_para = contact_cell.paragraphs[0]
        contact_para.add_run('YOUR CONTACT').bold = True
        contact_para.add_run('\n\nat e-aces.com\n\nPhone:\n\nMail:')
        table2.rows[row_idx].cells[1].text = ''
        row_idx += 1
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Professional Experience heading (centered, bold)
        exp_heading = doc.add_heading('Professional experience', 1)
        exp_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Professional experiences in tables
        for exp in expert.get('professional_experiences', []):
            # Create table for each experience
            exp_table = doc.add_table(rows=2, cols=2)
            exp_table.style = 'Table Grid'
            exp_table.columns[0].width = Inches(1.5)
            exp_table.columns[1].width = Inches(5.0)
            
            # First row: Period and Company/Role/Details
            period_cell = exp_table.rows[0].cells[0]
            period_para = period_cell.paragraphs[0]
            period_para.add_run(f"{exp.get('from', '')} -\n{exp.get('to', '')}").bold = True
            
            details_cell = exp_table.rows[0].cells[1]
            details_para = details_cell.paragraphs[0]
            
            # Company name (bold)
            details_para.add_run(exp.get('company', '')).bold = True
            details_para.add_run(f"\n{exp.get('role', '')}\n\n")
            
            # Tasks (bullet points)
            if exp.get('tasks'):
                for task in exp.get('tasks', []):
                    task_para = details_cell.add_paragraph(task, style='List Bullet')
                
                details_cell.add_paragraph()
            
            # Technologies
            if exp.get('technologies'):
                tech_para = details_cell.add_paragraph()
                tech_para.add_run('Technologies:\n').bold = False
                tech_para.add_run(', '.join(exp.get('technologies', [])))
                details_cell.add_paragraph()
            
            # Projects
            if exp.get('projects'):
                proj_para = details_cell.add_paragraph()
                proj_run = proj_para.add_run('Projects:')
                proj_run.underline = True
                
                for proj in exp.get('projects', []):
                    proj_item = details_cell.add_paragraph(style='List Bullet')
                    proj_item.add_run(f"{proj.get('name', '')}\n").bold = True
                    if proj.get('role'):
                        proj_item.add_run(f"Role: {proj.get('role', '')}\n").italic = True
                    if proj.get('description'):
                        proj_item.add_run(f"{proj.get('description', '')}\n").bold = True
                    if proj.get('responsibilities'):
                        proj_item.add_run(f"{proj.get('responsibilities', '')}")
            
            # Second row (empty separator)
            exp_table.rows[1].cells[0].text = ''
            exp_table.rows[1].cells[1].text = ''
            
            doc.add_paragraph()
        
        # Save to bytes
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        print("Word document generated successfully.")

        return send_file(
            file_stream,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name='expert_profile.docx'
        )
        
    except Exception as e:
        print(f"Error in generate_word: {str(e)}")
        return jsonify({'error': str(e)}), 500

def add_table_row_bold(label_cell, label, value, value_cell):
    """Add a row with bold label and normal value"""
    label_para = label_cell.paragraphs[0]
    label_para.add_run(label).bold = True
    value_cell.text = value

def add_table_section_bold(cell, text):
    """Add bold text to a cell"""
    para = cell.paragraphs[0]
    para.add_run(text).bold = True

if __name__ == '__main__':
    app.run(debug=True, port=5000)
